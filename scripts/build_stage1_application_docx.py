from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "stage1_application.docx"


def set_run_font(run, name="Microsoft YaHei", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=11)
    set_paragraph_spacing(p, before=0, after=6, line=1.15)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color="1F1F1F")
    set_paragraph_spacing(p, before=16, after=8, line=1.1)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=11)
    set_paragraph_spacing(p, before=0, after=4, line=1.15)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Moon Mustache 项目申报书")
    set_run_font(run, size=20, bold=True, color="000000")
    set_paragraph_spacing(p, before=0, after=12, line=1.0)


def add_metadata(doc):
    lines = [
        "项目名称：Moon Mustache：Mustache 模板引擎的 MoonBit 实现",
        "参赛者：[填写姓名]",
        "联系方式：[填写手机号 / 邮箱]",
        "GitHub 仓库链接：[填写你自己的 GitHub 仓库链接]",
        "Gitlink 仓库链接：[填写你自己的 Gitlink 仓库链接]",
        "项目方向：MoonBit 模板渲染基础库 / 工程基础设施",
        "是否为移植项目：是",
    ]
    for line in lines:
        add_body_paragraph(doc, line)


def build_doc():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_heading(doc, "基本信息")
    add_metadata(doc)

    add_heading(doc, "项目简介")
    add_body_paragraph(
        doc,
        "Moon Mustache 计划将 Mustache 模板渲染能力引入 MoonBit 生态，为脚手架生成、代码生成、配置文件渲染、"
        "静态内容拼装、邮件模板和消息正文生成等场景提供一个轻量、稳定、可复用的模板引擎。项目面向需要在 MoonBit 中"
        "生成字符串输出、构建可复用文本模板能力的库作者、工具开发者和应用开发者，提供模板解析、上下文绑定、渲染执行、"
        "Partial 组合、转义输出以及文件接口和命令行工具等能力。"
    )
    add_body_paragraph(
        doc,
        "本项目的目标不是扩展一个功能过重的模板语言，而是优先补齐 MoonBit 生态当前缺少的基础模板能力，形成一个边界清楚、"
        "接口稳定、文档完整、测试可运行、能够被其他项目直接复用的开源基础库。"
    )

    add_heading(doc, "核心功能范围")
    add_bullet(doc, "提供 Mustache 模板的词法分析、语法解析和中间表示，便于后续渲染、调试和扩展。")
    add_bullet(doc, "支持变量插值、Section、Inverted Section、Comment、Partial、Set Delimiter 等核心语法。")
    add_bullet(doc, "提供上下文栈、路径查找和值解析机制，支持常见数据结构绑定。")
    add_bullet(doc, "提供 HTML escaping 与原样输出能力，覆盖模板渲染中的常见安全需求。")
    add_bullet(doc, "提供 `parse_template`、`render_string`、`render_file` 等基础接口。")
    add_bullet(doc, "提供简单 CLI 工具，支持模板文件与数据文件输入、渲染结果输出。")
    add_bullet(doc, "提供模板错误定位、基础调试信息和典型异常场景测试。")
    add_bullet(doc, "提供与 mustache/spec 对照的兼容性测试、示例和迁移记录。")
    add_bullet(doc, "提供 README 示例，覆盖字符串渲染、文件渲染、Partial 使用和数据绑定场景。")
    add_bullet(doc, "提供持续集成配置，覆盖检查、构建和测试流程。")

    add_heading(doc, "移植或参考说明")
    add_bullet(doc, "原项目名称：mustache.js")
    add_bullet(doc, "原项目链接：https://github.com/janl/mustache.js")
    add_bullet(doc, "参考规范：https://github.com/mustache/spec")
    add_bullet(doc, "原项目许可证：MIT License")
    add_bullet(doc, "本项目许可证：MIT License")

    add_heading(doc, "与原项目相比，本项目会做以下简化和重新设计")
    add_bullet(doc, "使用 MoonBit 原生包结构、类型系统和测试方式组织代码，而不是复刻 JavaScript 的运行时接口和工程组织。")
    add_bullet(doc, "优先实现适合 MoonBit 工程场景的核心模板能力，弱化浏览器、DOM 或 Node.js 运行环境相关假设。")
    add_bullet(doc, "以可嵌入库和命令行工具作为主要交付形式，方便接入 MoonBit CLI、脚手架和自动化流程。")
    add_bullet(doc, "对数据模型、错误处理、上下文查找和 API 形态进行 MoonBit 化重写，而不是逐层照搬原实现。")
    add_bullet(doc, "以规范兼容测试、示例、README 和工程可复用性为主要交付目标，方便后续发布到 mooncakes.io 并服务其他项目。")

    add_heading(doc, "项目实施计划")
    add_bullet(doc, "第一阶段完成仓库初始化、基础数据结构、解析器原型、最小渲染链路、示例与早期测试，并形成 10-20 次有效提交。")
    add_bullet(doc, "第二阶段补全核心语法支持、Partial、文件接口、CLI、README、持续集成和回归测试。")
    add_bullet(doc, "最终交付公开可复用的 MoonBit 模板渲染库、命令行工具、测试集、示例和开发文档，并满足基础验收要求。")

    add_heading(doc, "项目规模预估")
    add_body_paragraph(
        doc,
        "预计主体实现约为 4k-6k 行有效 MoonBit 代码，另包含测试、示例、文档和 CI 配置，整体规模控制在赛事建议区间内。"
        "功能范围会优先聚焦 Mustache 规范核心能力与工程可用性，不扩展到超出 Mustache 范畴的大量逻辑模板特性，"
        "以保证项目目标聚焦、实现完整、验收明确。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
